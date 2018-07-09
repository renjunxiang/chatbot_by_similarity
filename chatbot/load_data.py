import re
import os
import docx

localpath = os.path.dirname(__file__)


def _docx2text(file_path=None):
    '''

    :param file_path: 文件路径
    :return:
    '''
    document = docx.Document(file_path)

    data_str = []
    for i in range(len(document.paragraphs)):
        if document.paragraphs[i].text != '':
            text_one = document.paragraphs[i].text
            data_str.append(text_one)

    return data_str


def load_data(type='chat'):
    if type == 'chat':
        f = open(localpath + '/data/xiaohuangji50w_nofenci.conv', encoding='utf-8')
        data = []
        line = True
        while line:
            line = f.readline()
            if line != 'E\n':
                re_func = re.compile('[M ,\n]')
                data.append(re_func.sub('', line))
        f.close()
        return data
    elif type == 'knowledge':
        texts = []
        folder_names = ['材料清单', '法律法规', '医疗器械']
        for folder_name in folder_names:
            folder_path = localpath + '/data/%s' % (folder_name)
            file_paths = os.listdir(folder_path)
            for file_path in file_paths:
                try:
                    text = _docx2text(file_path=folder_path + '/' + file_path)
                    texts += text
                except Exception as e:
                    print('错误：' + folder_path + '/' + file_path)
                    texts.append('')
        return texts
