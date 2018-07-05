from chatbot import chatbot
import docx
import os
import pickle


def docx2text(file_path=None):
    '''

    :param file_path: 文件路径
    :return:
    '''
    document = docx.Document(file_path)

    data_str = []
    for i in range(len(document.paragraphs)):
        if document.paragraphs[i].text != '':
            text_one = document.paragraphs[i].text
            data_str.append(text_one)

    return data_str


texts = []
folder_names = ['材料清单', '法律法规', '医疗器械']
for folder_name in folder_names:
    folder_path = 'D:/work/svn/wonders_nlp/data/%s' % (folder_name)
    file_paths = os.listdir(folder_path)
    for file_path in file_paths:
        try:
            text = docx2text(file_path=folder_path + '/' + file_path)
            texts+=text
        except Exception as e:
            print('错误：' + folder_path + '/' + file_path)
            texts.append('')

chatbot_try = chatbot()
chatbot_try.train(texts=texts)
with open('./model_word_document.pkl',mode='wb') as f:
    pickle.dump(chatbot_try,f)


